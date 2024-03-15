#python3
#Write a program that would generate a Word document with custom invitations

import docx #pip install docx
from docx.enum.text import WD_BREAK #pip install python-docx

doc = docx.Document('invitation.docx') #create the empty document in the same folder with
                                       #different fonts and save their names as you'll need them later

names = open('guests.txt') #download the file from https://nostarch.com /automatestuff2/
names = names.read()
names = names.split(sep='\n')

for name in names:
    doc.add_paragraph('It would be a pleasure to have the company of').style = 'inviteStyleLine1'
    doc.add_paragraph(name).style = 'Name'
    doc.add_paragraph('at 11010 Memory Lane on the Evening of').style = 'inviteStyleLine2'
    doc.add_paragraph('April 1st').style = 'invitedate'
    date_line = doc.add_paragraph('at 7 o\'clock')
    date_line.runs[0].add_break(break_type=WD_BREAK.PAGE)
    date_line.style = 'inviteStyleLine2'

doc.save('invitations.docx')
